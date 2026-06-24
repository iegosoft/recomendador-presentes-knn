"""Gera o documento .docx do relatório do projeto Recomendador de Presentes."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

CAMINHO_SAIDA = "documentacao/relatorio-recomendador.docx"

COR_ROXO = RGBColor(0x6C, 0x2B, 0xD9)
COR_CINZA_ESCURO = RGBColor(0x15, 0x10, 0x1F)
COR_TEXTO = RGBColor(0x1A, 0x1A, 0x1A)
COR_SUAVE = RGBColor(0x6F, 0x6A, 0x78)
COR_CODIGO_FUNDO = RGBColor(0xF3, 0xF0, 0xFA)


def set_cell_background(cell, color_hex: str):
    """Define cor de fundo de uma célula de tabela."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tcPr.append(shd)


def add_heading(doc: Document, text: str, level: int = 1):
    """Adiciona um título com cor e fonte customizados."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18 if level == 1 else 12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Calibri"
    run.font.color.rgb = COR_ROXO if level <= 2 else COR_CINZA_ESCURO
    run.font.size = Pt(20 if level == 1 else 15 if level == 2 else 13)
    return p


def add_body(doc: Document, text: str, bold_parts: list[str] | None = None):
    """Adiciona parágrafo de corpo com suporte a partes em negrito."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.space_before = Pt(2)

    if not bold_parts:
        run = p.add_run(text)
        run.font.name = "Calibri"
        run.font.size = Pt(11)
        run.font.color.rgb = COR_TEXTO
        return p

    remaining = text
    for bold in bold_parts:
        parts = remaining.split(bold, 1)
        if len(parts) == 2:
            if parts[0]:
                r = p.add_run(parts[0])
                r.font.name = "Calibri"
                r.font.size = Pt(11)
                r.font.color.rgb = COR_TEXTO
            rb = p.add_run(bold)
            rb.bold = True
            rb.font.name = "Calibri"
            rb.font.size = Pt(11)
            rb.font.color.rgb = COR_ROXO
            remaining = parts[1]
    if remaining:
        r = p.add_run(remaining)
        r.font.name = "Calibri"
        r.font.size = Pt(11)
        r.font.color.rgb = COR_TEXTO
    return p


def add_code(doc: Document, code: str):
    """Adiciona bloco de código com fundo colorido."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.left_indent = Cm(0.8)
    run = p.add_run(code)
    run.font.name = "Courier New"
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x3A, 0x00, 0x9C)
    # fundo via shading no parágrafo
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F3F0FA")
    pPr.append(shd)
    return p


def add_bullet(doc: Document, text: str):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.font.color.rgb = COR_TEXTO
    return p


def add_placeholder_imagem(doc: Document, descricao: str):
    """Adiciona uma caixa de texto visual como espaço reservado para imagem."""
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    cell.width = Inches(6)
    set_cell_background(cell, "F3F0FA")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(30)
    p.paragraph_format.space_after = Pt(30)
    run = p.add_run(f"[ ESPAÇO RESERVADO PARA CAPTURA DE TELA ]\n{descricao}")
    run.font.name = "Calibri"
    run.font.size = Pt(10)
    run.font.color.rgb = COR_SUAVE
    run.italic = True
    doc.add_paragraph()


def gerar():
    doc = Document()

    # ── Margens ──────────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2)

    # ════════════════════════════════════════════════════════════════════════
    # 1. CAPA
    # ════════════════════════════════════════════════════════════════════════
    doc.add_paragraph()
    doc.add_paragraph()

    p_titulo = doc.add_paragraph()
    p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_titulo.add_run("Sistema de Recomendação de Presentes")
    r.bold = True
    r.font.name = "Calibri"
    r.font.size = Pt(26)
    r.font.color.rgb = COR_ROXO

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p_sub.add_run(
        "Baseado em Conteúdo com Similaridade de Cosseno\ne K-Vizinhos Mais Próximos (KNN)"
    )
    r2.font.name = "Calibri"
    r2.font.size = Pt(15)
    r2.font.color.rgb = COR_SUAVE

    doc.add_paragraph()

    p_disc = doc.add_paragraph()
    p_disc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p_disc.add_run("Disciplina de Álgebra Linear  ·  Prof. Fernando Soares")
    r3.font.name = "Calibri"
    r3.font.size = Pt(12)
    r3.font.color.rgb = COR_CINZA_ESCURO

    doc.add_paragraph()

    discentes = [
        "Ariane Nascimento Andrade",
        "Josielle Santos da Silva",
        "Juliane Vitória de Souza Silva",
        "Mauricio Carvalho de Castro",
        "Iêgo Sérgio Costa de Souza",
        "Rayssa Oliveira Martins das Chagas",
    ]
    for nome in discentes:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(nome)
        r.font.name = "Calibri"
        r.font.size = Pt(12)
        r.font.color.rgb = COR_TEXTO

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    # 2. INTRODUÇÃO
    # ════════════════════════════════════════════════════════════════════════
    add_heading(doc, "1. Introdução")

    add_body(
        doc,
        "O Recomendador de Presentes é um sistema web que sugere presentes personalizados a partir "
        "do perfil de quem vai receber — informando idade, gênero, orçamento, ocasião e interesses. "
        "O problema que ele resolve é simples de enunciar e difícil na prática: como escolher um "
        "presente relevante para alguém sem conhecer seus gostos em detalhes?",
    )
    add_body(
        doc,
        "A abordagem adotada é a Filtragem Baseada em Conteúdo (Content-Based Filtering): em vez "
        "de comparar o comportamento de outros usuários, o sistema compara matematicamente os atributos "
        "de cada item do catálogo com o perfil informado. Isso é possível porque tanto o perfil quanto "
        "os itens são representados como vetores em um espaço multidimensional — e a distância entre "
        "esses vetores é medida com a Similaridade de Cosseno, conceito central de Álgebra Linear.",
        bold_parts=["Filtragem Baseada em Conteúdo", "Similaridade de Cosseno"],
    )

    # ════════════════════════════════════════════════════════════════════════
    # 3. TECNOLOGIAS
    # ════════════════════════════════════════════════════════════════════════
    add_heading(doc, "2. Tecnologias Utilizadas")

    tech_itens = [
        ("Python 3.14", "Linguagem principal — backend e scripts de dados"),
        ("Flask 3.1.3", "Framework web: serve a interface HTML e expõe a API REST"),
        ("scikit-learn 1.9.0", "Implementação do KNN: NearestNeighbors(metric='cosine')"),
        ("NumPy 2.5.0", "Vetores multi-hot e cálculo do produto interno (np.dot)"),
        ("pandas 3.0.3", "Carregamento e filtragem do catálogo CSV"),
        ("pytest 8.4.2", "24 testes automatizados (integração + unitários)"),
        ("GitHub Actions", "CI/CD: testes rodam automaticamente em cada pull request"),
    ]

    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for cell, txt in zip(hdr, ["Tecnologia / Versão", "Uso no Projeto"]):
        set_cell_background(cell, "6C2BD9")
        p = cell.paragraphs[0]
        run = p.add_run(txt)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.name = "Calibri"
        run.font.size = Pt(10)

    for tech, uso in tech_itens:
        row = table.add_row().cells
        for cell, txt in zip(row, [tech, uso]):
            p = cell.paragraphs[0]
            run = p.add_run(txt)
            run.font.name = "Calibri"
            run.font.size = Pt(10)

    doc.add_paragraph()
    add_heading(doc, "Estrutura do Projeto", level=3)
    add_body(doc, "O projeto é organizado em módulos com responsabilidades bem definidas:")
    estrutura = [
        "app.py — servidor Flask, validação de entrada e rotas da API",
        "recommender.py — motor de recomendação (KNN, cosseno, filtros, diversidade)",
        "data/presentes.csv — catálogo de 1.512 itens",
        "templates/index.html + static/ — interface web (HTML, CSS, JS)",
        "tests/ — 24 testes automatizados",
        ".github/workflows/testes.yml — pipeline de CI com GitHub Actions",
    ]
    for item in estrutura:
        add_bullet(doc, item)

    # ════════════════════════════════════════════════════════════════════════
    # 4. FUNDAMENTAÇÃO TÉCNICA
    # ════════════════════════════════════════════════════════════════════════
    add_heading(doc, "3. Fundamentação Técnica")

    # 4.1
    add_heading(doc, "3.1  Representação dos Dados como Vetores", level=2)
    add_body(
        doc,
        "Para que seja possível calcular a similaridade entre um perfil e um item do catálogo, "
        "ambos precisam estar representados na mesma linguagem matemática: vetores.",
    )
    add_body(
        doc,
        "O projeto define um vocabulário fixo de 19 tags (interesses possíveis): tecnologia, fitness, "
        "viagem, música, fotografia, arte, jogos, leitura, culinária, moda, beleza, bem-estar, esportes, "
        "café, vinho, artesanato, decoração, jardinagem e pets. Cada tag ocupa uma posição fixa nesse "
        "vocabulário.",
        bold_parts=["19 tags"],
    )
    add_body(
        doc,
        "Cada item ou perfil é codificado como um vetor binário de 19 dimensões — 1 se aquela tag "
        "está presente, 0 caso contrário. Essa técnica chama-se multi-hot encoding:",
        bold_parts=["multi-hot encoding"],
    )
    add_code(
        doc,
        'Exemplo — "Tapete de yoga"  (tags: fitness, bem-estar, esportes)\n'
        "v_item   = [0, 1, 0, 0, 0, 0, 0, 0, 0, 0, 0, 1, 1, 0, 0, 0, 0, 0, 0]\n"
        "                ^fitness               ^bem-estar ^esportes\n\n"
        "Perfil do usuário  (interesses: fitness, esportes)\n"
        "v_perfil = [0, 1, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 1, 0, 0, 0, 0, 0, 0]",
    )
    add_body(
        doc,
        "Com isso, tanto o perfil quanto cada item existem no mesmo espaço vetorial ℝ¹⁹, "
        "o que permite aplicar diretamente as operações de Álgebra Linear para medir a proximidade entre eles.",
        bold_parts=["ℝ¹⁹"],
    )

    # 4.2
    add_heading(doc, "3.2  Similaridade de Cosseno", level=2)
    add_body(
        doc,
        "A Similaridade de Cosseno mede o ângulo entre dois vetores no espaço multidimensional. "
        "Ela responde: esses dois vetores apontam na mesma direção?",
        bold_parts=["ângulo entre dois vetores"],
    )
    add_body(doc, "A fórmula é derivada da definição de produto interno:")
    add_code(
        doc,
        "         a · b              Σ (aᵢ × bᵢ)\n"
        "cos(θ) = ──────── =  ─────────────────────────\n"
        "         ‖a‖ × ‖b‖     √(Σ aᵢ²) × √(Σ bᵢ²)",
    )
    add_body(
        doc,
        "O resultado vai de 0 (vetores perpendiculares — nenhuma similaridade) a 1 (vetores paralelos — "
        "similaridade total). Para vetores binários como os deste projeto, o numerador a · b conta "
        "quantas tags os dois têm em comum; o denominador normaliza pelo tamanho de cada conjunto.",
        bold_parts=["a · b"],
    )
    add_body(
        doc,
        "Por que cosseno e não distância euclidiana? A distância euclidiana é sensível à magnitude "
        "(comprimento) dos vetores: dois itens com muitas tags em comum mas também com muitas tags "
        "diferentes teriam distância grande, mesmo sendo semanticamente próximos. O cosseno captura "
        "apenas a direção (interesses compartilhados) e ignora a magnitude — exatamente o que se "
        "precisa para comparar interesses.",
        bold_parts=["Por que cosseno e não distância euclidiana?"],
    )
    add_body(doc, "Como aparece no código:")
    add_code(
        doc,
        "from sklearn.neighbors import NearestNeighbors\n\n"
        "# Cada linha é o vetor multi-hot de um item do catálogo\n"
        "modelo = NearestNeighbors(metric='cosine', n_neighbors=tamanho_pool)\n"
        "modelo.fit(matriz_candidatos)   # shape: (N_itens, 19)\n"
        "_, posicoes = modelo.kneighbors([perfil_tags])  # shape: (19,)",
    )

    # 4.3
    add_heading(doc, "3.3  K-Vizinhos Mais Próximos (KNN)", level=2)
    add_body(
        doc,
        "O KNN é um algoritmo não-paramétrico: em vez de criar um modelo treinado com pesos, "
        "ele procura diretamente no espaço vetorial quais pontos (itens) estão mais próximos "
        "de um ponto de consulta (perfil do usuário).",
        bold_parts=["KNN"],
    )
    add_body(doc, "No pipeline deste projeto, o KNN atua da seguinte forma:")
    add_bullet(doc, "Os itens do catálogo (após filtros de idade, gênero e orçamento) formam a matriz de pontos no ℝ¹⁹")
    add_bullet(doc, "O vetor do perfil é o ponto de consulta")
    add_bullet(doc, "A métrica de distância é cosseno: d = 1 – cos(θ), de modo que menor distância = maior similaridade")
    add_bullet(doc, "O KNN retorna os K itens com menor distância — o pool de candidatos para avaliação final")
    doc.add_paragraph()
    add_body(
        doc,
        "O parâmetro K é definido como top_n × 20 (onde top_n = 6 recomendações desejadas). "
        "Um K grande é necessário porque o catálogo tem 21 variantes por produto-base (3 estilos × 7 "
        "faixas de preço) com tags idênticas — um K pequeno ficaria tomado pelas variantes do mesmo "
        "produto.",
        bold_parts=["K = top_n × 20"],
    )
    add_body(
        doc,
        "Como os dois conceitos se encaixam: o KNN usa a distância cosseno para encontrar os candidatos "
        "mais próximos do perfil. Depois, o Coeficiente de Sobreposição (overlap / min(|perfil|, |item|)) "
        "calcula a compatibilidade percentual exibida ao usuário — uma métrica que resolve o problema "
        "de resultados vazios quando o usuário seleciona interesses de muitas categorias diferentes.",
        bold_parts=["Coeficiente de Sobreposição", "overlap / min(|perfil|, |item|)"],
    )

    # ════════════════════════════════════════════════════════════════════════
    # 5. FLUXO DO SISTEMA
    # ════════════════════════════════════════════════════════════════════════
    add_heading(doc, "4. Fluxo do Sistema")

    add_code(
        doc,
        "Usuário preenche o formulário\n"
        "  ↓\n"
        "Filtros rígidos  →  elimina itens fora da faixa etária, gênero e orçamento\n"
        "  ↓\n"
        "Vetorização  →  converte interesses em vetor multi-hot de 19 dimensões\n"
        "  ↓\n"
        "KNN com métrica cosseno  →  encontra os K itens mais próximos do perfil\n"
        "  ↓\n"
        "Coeficiente de Sobreposição  →  calcula % de compatibilidade de cada candidato\n"
        "  ↓\n"
        "Piso mínimo 58%  →  descarta itens com baixa afinidade\n"
        "  ↓\n"
        "Deduplicação + Diversidade  →  no máximo 2 por categoria, sem variantes repetidas\n"
        "  ↓\n"
        "Resultado: até 6 presentes ordenados por compatibilidade",
    )

    # ════════════════════════════════════════════════════════════════════════
    # 6. TELAS DO SISTEMA
    # ════════════════════════════════════════════════════════════════════════
    add_heading(doc, "5. Telas do Sistema")

    add_heading(doc, "5.1  Formulário de Perfil", level=2)
    add_body(
        doc,
        "Tela inicial onde o usuário informa a idade, gênero, orçamento, ocasião e seleciona "
        "os interesses do destinatário por meio de chips clicáveis.",
    )
    add_placeholder_imagem(doc, "Tela inicial — formulário de perfil com chips de interesses")

    add_heading(doc, "5.2  Cards de Resultado", level=2)
    add_body(
        doc,
        "Após o envio, o sistema exibe até 6 presentes recomendados em cards com gradiente "
        "por categoria, nome do produto, preço e percentual de compatibilidade.",
    )
    add_placeholder_imagem(doc, "Tela de resultados — grid de cards com % de compatibilidade")

    add_heading(doc, "5.3  Aba Sobre o Projeto", level=2)
    add_body(
        doc,
        "Aba acessível no topo da interface com as informações da disciplina, professor e "
        "equipe de desenvolvimento.",
    )
    add_placeholder_imagem(doc, "Aba 'Sobre o Projeto' — disciplina, professor e equipe")

    # ════════════════════════════════════════════════════════════════════════
    # 7. CONSIDERAÇÕES FINAIS
    # ════════════════════════════════════════════════════════════════════════
    add_heading(doc, "6. Considerações Finais")

    add_heading(doc, "O que funcionou bem", level=3)
    add_bullet(doc, "A representação vetorial multi-hot permitiu aplicar diretamente operações de Álgebra Linear para comparar perfis e itens")
    add_bullet(doc, "O KNN com métrica cosseno provou ser eficiente e interpretável: basta calcular o ângulo entre vetores para medir afinidade")
    add_bullet(doc, "O Coeficiente de Sobreposição resolveu o problema de resultados vazios para usuários com interesses diversificados — uma limitação das métricas mais simples")
    add_bullet(doc, "Os 24 testes automatizados e o pipeline de CI garantiram que nenhuma alteração quebrasse o sistema silenciosamente")

    add_heading(doc, "Limitações", level=3)
    add_bullet(doc, "O catálogo é estático (1.512 itens pré-definidos) — não há integração com lojas reais nem atualização automática de preços")
    add_bullet(doc, "O sistema não possui memória entre sessões: cada recomendação parte do zero, sem histórico do usuário")
    add_bullet(doc, "A abordagem é puramente baseada em conteúdo — não considera o que usuários semelhantes compraram (filtragem colaborativa)")

    add_heading(doc, "Próximos Passos", level=3)
    add_bullet(doc, "Integrar com APIs de lojas (Amazon, Mercado Livre) para catálogo dinâmico e links de compra")
    add_bullet(doc, "Adicionar perfil persistente e histórico de recomendações por usuário")
    add_bullet(doc, "Explorar abordagem híbrida combinando filtragem baseada em conteúdo com filtragem colaborativa")
    add_bullet(doc, "Implementar feedback explícito ('gostei / não gostei') para refinar as recomendações ao longo do tempo")

    # ────────────────────────────────────────────────────────────────────────
    os.makedirs("documentacao", exist_ok=True)
    doc.save(CAMINHO_SAIDA)
    print(f"Documento gerado: {CAMINHO_SAIDA}")


if __name__ == "__main__":
    gerar()
